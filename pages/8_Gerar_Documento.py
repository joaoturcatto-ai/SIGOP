import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from utils.db import fetch_table

st.set_page_config(page_title="Gerar Documento - SIGOP", page_icon="📄", layout="wide")
st.title("📄 Gerar Ordem de Operação")
st.caption("Gera o documento oficial da operação em Word (.docx) ou PDF")


def formatar_datas_folga(raw):
    """Converte a string de datas separadas por vírgula em uma lista de datas
    formatadas (dd/mm/aaaa), ignorando qualquer valor inválido."""
    if not raw:
        return []
    datas_fmt = []
    for d in str(raw).split(","):
        d = d.strip()
        if not d or d.lower() in ("nan", "none", "nat"):
            continue
        dt = pd.to_datetime(d, errors="coerce")
        if pd.notna(dt):
            datas_fmt.append(dt.strftime("%d/%m/%Y"))
    return datas_fmt


def placa_disponivel(row):
    p_of = row.get("placa_oficial")
    if pd.notna(p_of) and str(p_of).strip():
        return str(p_of)
    p_res = row.get("placa_reservada")
    if pd.notna(p_res) and str(p_res).strip():
        return str(p_res)
    return "Sem Placa"


def nome_membro_equipe(row_m, mapa_servidores_local):
    """Retorna o nome de exibição de um membro da equipe, seja ele um
    servidor cadastrado no sistema ou uma pessoa externa (sem servidor_id)."""
    sid = row_m.get("servidor_id")
    if pd.notna(sid):
        return mapa_servidores_local.get(int(sid), "Não encontrado")
    nome_ext = row_m.get("nome_externo")
    if pd.notna(nome_ext) and str(nome_ext).strip():
        return f"{nome_ext} (Externo)"
    return "Pessoa externa (sem nome)"


try:
    operacoes = fetch_table("operacoes", order_by="data_inicio")
    servidores = fetch_table("servidores")
    viaturas = fetch_table("viaturas")
    participantes = fetch_table("equipes_operacoes")
except Exception as e:
    st.error("⚠️ Erro ao carregar dados.")
    st.code(str(e), language="python")
    st.stop()

if operacoes.empty:
    st.info("Cadastre uma operação primeiro na página 'Operações'.")
    st.stop()

mapa_servidores_doc = {row["id"]: row["nome"] for _, row in servidores.iterrows()} if not servidores.empty else {}
mapa_cargo_doc = {row["id"]: row.get("cargo", "") for _, row in servidores.iterrows()} if not servidores.empty else {}
mapa_viaturas_doc = (
    {row["id"]: f"{row['identificacao']} — {placa_disponivel(row)}" for _, row in viaturas.iterrows()}
    if not viaturas.empty
    else {}
)

operacao_id = st.selectbox(
    "Selecione a operação",
    options=operacoes["id"].tolist(),
    format_func=lambda x: f"#{x} — {operacoes[operacoes['id'] == x]['nome'].values[0]}",
)

dados_op = operacoes[operacoes["id"] == operacao_id].iloc[0]
participantes_op = (
    participantes[participantes["operacao_id"] == operacao_id]
    if not participantes.empty
    else pd.DataFrame()
)

tipo_raw = dados_op.get("tipo")
tipo_op = str(tipo_raw) if pd.notna(tipo_raw) else "Operação"

delegado_id_op = dados_op.get("delegado_id")
delegado_nome = mapa_servidores_doc.get(delegado_id_op, "Não informado") if pd.notna(delegado_id_op) else "Não informado"

data_ini_fmt = (
    pd.to_datetime(dados_op["data_inicio"]).strftime("%d/%m/%Y") if pd.notna(dados_op.get("data_inicio")) else "Não definida"
)
data_fim_fmt = (
    pd.to_datetime(dados_op["data_fim"]).strftime("%d/%m/%Y") if pd.notna(dados_op.get("data_fim")) else "Não definida"
)
horario_fim_previsto = dados_op.get("horario_fim_previsto")
horario_completo = str(dados_op.get("horario", "—")) + (
    f" até {horario_fim_previsto} (previsto)" if pd.notna(horario_fim_previsto) else ""
)

# Monta a lista de equipes com seus membros, líder, viatura e status de folga
equipes_doc = []
if not participantes_op.empty:
    for nome_equipe in sorted(participantes_op["nome_equipe"].dropna().unique()):
        membros_eq = participantes_op[participantes_op["nome_equipe"] == nome_equipe]

        lider_membros = membros_eq[membros_eq.get("is_lider", False) == True]
        lider_nome = (
            nome_membro_equipe(lider_membros.iloc[0], mapa_servidores_doc)
            if not lider_membros.empty
            else "Não definido"
        )

        vtr_id = None
        for _, r in membros_eq.iterrows():
            if pd.notna(r.get("viatura_id")):
                vtr_id = r["viatura_id"]
                break
        viatura_txt = mapa_viaturas_doc.get(vtr_id, "Sem viatura designada")

        membros_lista = []
        for _, r in membros_eq.iterrows():
            nome_m = nome_membro_equipe(r, mapa_servidores_doc)
            cargo_m = mapa_cargo_doc.get(r["servidor_id"], "") if pd.notna(r.get("servidor_id")) else "Externo"
            marcador = " [LÍDER]" if r.get("is_lider", False) else ""

            if r.get("possui_folga", False):
                datas_f = formatar_datas_folga(r.get("folga_data"))
                dur_raw = r.get("folga_duracao")
                dur_txt = str(dur_raw) if pd.notna(dur_raw) else "Integral"
                folga_txt = (
                    f"Folga {dur_txt} em {', '.join(datas_f)}"
                    if datas_f
                    else f"Direito a folga {dur_txt} — data a definir"
                )
            else:
                folga_txt = "Sem folga"

            membros_lista.append(
                {"nome": nome_m, "cargo": cargo_m, "marcador": marcador, "folga": folga_txt}
            )

        equipes_doc.append(
            {
                "nome_equipe": nome_equipe,
                "lider": lider_nome,
                "viatura": viatura_txt,
                "membros": membros_lista,
            }
        )

st.markdown("---")
st.subheader("Pré-visualização")
st.write(f"**Tipo:** {tipo_op}")
st.write(f"**Operação:** {dados_op['nome']}")
st.write(f"**Período:** {data_ini_fmt} até {data_fim_fmt} — **Horário:** {horario_completo}")
st.write(f"**Local:** {dados_op.get('local', '—')} — **Cidade:** {dados_op.get('cidade', '—')}")
st.write(f"**Delegado responsável:** {delegado_nome}")
st.write("**Objetivo:**")
st.write(dados_op.get("objetivo") or "Não informado.")
st.write("**Briefing:**")
st.write(dados_op.get("briefing") or "Não informado.")

if equipes_doc:
    st.write("**Equipes escaladas:**")
    for eq in equipes_doc:
        st.markdown(f"**{eq['nome_equipe']}** — 👑 Líder: {eq['lider']} — 🚗 Viatura: {eq['viatura']}")
        df_eq = pd.DataFrame(
            [{"Nome": f"{m['nome']}{m['marcador']}", "Cargo": m["cargo"], "Folga": m["folga"]} for m in eq["membros"]]
        )
        df_eq.index = df_eq.index + 1
        st.dataframe(df_eq, use_container_width=True)
else:
    st.info("Nenhuma equipe escalada nesta operação ainda. Vá em 'Operações' para montar as equipes.")


def gerar_docx() -> BytesIO:
    doc = Document()

    titulo = doc.add_paragraph()
    run = titulo.add_run("POLÍCIA CIVIL DO ESTADO DE MATO GROSSO")
    run.bold = True
    run.font.size = Pt(14)
    titulo.alignment = 1

    subtitulo = doc.add_paragraph()
    run = subtitulo.add_run("Delegacia Especializada de Estelionato de Cuiabá/MT")
    run.font.size = Pt(11)
    subtitulo.alignment = 1

    doc.add_paragraph()
    nome_op = doc.add_paragraph()
    run = nome_op.add_run(f"{tipo_op.upper()} — {dados_op['nome'].upper()}")
    run.bold = True
    run.font.size = Pt(13)
    nome_op.alignment = 1

    doc.add_paragraph()
    doc.add_paragraph(f"Período: {data_ini_fmt} até {data_fim_fmt}    Horário: {horario_completo}")
    doc.add_paragraph(f"Local: {dados_op.get('local', '—')}    Cidade: {dados_op.get('cidade', '—')}")
    doc.add_paragraph(f"Delegado responsável: {delegado_nome}")

    doc.add_heading("Objetivo", level=2)
    doc.add_paragraph(dados_op.get("objetivo") or "Não informado.")

    doc.add_heading("Briefing", level=2)
    doc.add_paragraph(dados_op.get("briefing") or "Não informado.")

    doc.add_heading("Equipes", level=2)
    if equipes_doc:
        for eq in equipes_doc:
            doc.add_heading(
                f"{eq['nome_equipe']} — Líder: {eq['lider']} — Viatura: {eq['viatura']}", level=3
            )
            tabela = doc.add_table(rows=1, cols=3)
            tabela.style = "Light Grid Accent 1"
            hdr = tabela.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Nome", "Cargo", "Folga"
            for m in eq["membros"]:
                cells = tabela.add_row().cells
                cells[0].text = f"{m['nome']}{m['marcador']}"
                cells[1].text = m["cargo"]
                cells[2].text = m["folga"]
    else:
        doc.add_paragraph("Nenhuma equipe escalada.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def gerar_pdf() -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    elementos = []

    elementos.append(Paragraph("<b>POLÍCIA CIVIL DO ESTADO DE MATO GROSSO</b>", styles["Title"]))
    elementos.append(Paragraph("Delegacia Especializada de Estelionato de Cuiabá/MT", styles["Normal"]))
    elementos.append(Spacer(1, 12))
    elementos.append(
        Paragraph(f"<b>{tipo_op.upper()} — {dados_op['nome'].upper()}</b>", styles["Heading2"])
    )
    elementos.append(Spacer(1, 8))

    elementos.append(
        Paragraph(
            f"<b>Período:</b> {data_ini_fmt} até {data_fim_fmt} &nbsp;&nbsp; <b>Horário:</b> {horario_completo}",
            styles["Normal"],
        )
    )
    elementos.append(
        Paragraph(
            f"<b>Local:</b> {dados_op.get('local', '—')} &nbsp;&nbsp; <b>Cidade:</b> {dados_op.get('cidade', '—')}",
            styles["Normal"],
        )
    )
    elementos.append(Paragraph(f"<b>Delegado responsável:</b> {delegado_nome}", styles["Normal"]))
    elementos.append(Spacer(1, 12))

    elementos.append(Paragraph("<b>Objetivo</b>", styles["Heading3"]))
    elementos.append(Paragraph(dados_op.get("objetivo") or "Não informado.", styles["Normal"]))
    elementos.append(Spacer(1, 8))

    elementos.append(Paragraph("<b>Briefing</b>", styles["Heading3"]))
    elementos.append(
        Paragraph((dados_op.get("briefing") or "Não informado.").replace("\n", "<br/>"), styles["Normal"])
    )
    elementos.append(Spacer(1, 12))

    elementos.append(Paragraph("<b>Equipes</b>", styles["Heading3"]))
    if equipes_doc:
        for eq in equipes_doc:
            elementos.append(Spacer(1, 6))
            elementos.append(
                Paragraph(
                    f"<b>{eq['nome_equipe']}</b> — Líder: {eq['lider']} — Viatura: {eq['viatura']}",
                    styles["Normal"],
                )
            )
            tabela_dados = [["Nome", "Cargo", "Folga"]]
            for m in eq["membros"]:
                tabela_dados.append([f"{m['nome']}{m['marcador']}", m["cargo"], m["folga"]])
            tabela = Table(tabela_dados, hAlign="LEFT")
            tabela.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ]
                )
            )
            elementos.append(tabela)
    else:
        elementos.append(Paragraph("Nenhuma equipe escalada.", styles["Normal"]))

    doc.build(elementos)
    buffer.seek(0)
    return buffer


st.markdown("---")
col1, col2 = st.columns(2)
with col1:
    if st.button("📝 Gerar documento Word (.docx)"):
        buffer = gerar_docx()
        st.download_button(
            "⬇️ Baixar Word",
            data=buffer,
            file_name=f"ordem_operacao_{dados_op['nome'].replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
with col2:
    if st.button("📄 Gerar documento PDF"):
        buffer = gerar_pdf()
        st.download_button(
            "⬇️ Baixar PDF",
            data=buffer,
            file_name=f"ordem_operacao_{dados_op['nome'].replace(' ', '_')}.pdf",
            mime="application/pdf",
        )
